from pathlib import Path

import docx
import pytest
from reportlab.pdfgen import canvas

from bot.parsers import extract_text


def test_txt_utf8_round_trip(tmp_path: Path) -> None:
    path = tmp_path / "sample.txt"
    original = "hello, мир, café"
    path.write_text(original, encoding="utf-8")

    assert extract_text(path) == original


def test_docx_extracts_text(tmp_path: Path) -> None:
    path = tmp_path / "sample.docx"
    document = docx.Document()
    document.add_paragraph("First line")
    document.add_paragraph("Вторая строка")
    document.save(path)

    text = extract_text(path)
    assert "First line" in text
    assert "Вторая строка" in text


def test_pdf_extracts_text(tmp_path: Path) -> None:
    path = tmp_path / "sample.pdf"
    c = canvas.Canvas(str(path))
    c.drawString(100, 750, "PDF hello")
    c.drawString(100, 730, "Привет PDF")
    c.save()

    text = extract_text(path)
    assert "PDF hello" in text


def test_unsupported_extension_raises(tmp_path: Path) -> None:
    path = tmp_path / "sample.md"
    path.write_text("content", encoding="utf-8")

    with pytest.raises(ValueError, match="unsupported file type"):
        extract_text(path)
